from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path
from typing import Any


def _office_cfg(context: dict[str, Any]) -> dict[str, Any]:
    return context.get("config", {}).get("office") or {}


def _driver(context: dict[str, Any]) -> str:
    return str(_office_cfg(context).get("driver") or os.environ.get("URISYS_OFFICE_DRIVER") or "mock")


def _output_dir(context: dict[str, Any]) -> Path:
    cfg = _office_cfg(context)
    root = Path(cfg.get("output_dir") or context.get("config", {}).get("office_output_dir") or "data/office")
    root.mkdir(parents=True, exist_ok=True)
    return root


def _doc_state(context: dict[str, Any]) -> dict[str, Any]:
    host = context.get("params", {}).get("host", "local")
    docs = context.setdefault("state", {}).setdefault("office_documents", {})
    return docs.setdefault(host, {"path": None, "content": "", "title": "untitled"})


def _real_allowed(context: dict[str, Any]) -> bool:
    return bool(context.get("allow_real") or os.environ.get("URISYS_ALLOW_REAL") == "1")


def status(payload: dict[str, Any], context: dict[str, Any]) -> dict[str, Any]:
    del payload
    doc = _doc_state(context)
    return {
        "driver": _driver(context),
        "document_path": doc.get("path"),
        "title": doc.get("title"),
        "supports": ["mock", "writer", "libreoffice"],
    }


def document_open(payload: dict[str, Any], context: dict[str, Any]) -> dict[str, Any]:
    path = payload.get("path") or payload.get("template")
    title = payload.get("title") or (Path(str(path)).stem if path else "untitled")
    driver = _driver(context)
    if context.get("dry_run"):
        return {"dry_run": True, "driver": driver, "path": path, "title": title}
    doc = _doc_state(context)
    out = _output_dir(context)
    if path:
        src = Path(str(path))
        if src.exists():
            dest = out / src.name
            shutil.copy2(src, dest)
            doc.update({"path": str(dest), "content": dest.read_text(encoding="utf-8", errors="replace"), "title": title})
            return {"opened": True, "path": str(dest), "driver": driver}
    doc.update({"path": str(out / f"{title}.txt"), "content": payload.get("content") or "", "title": title})
    Path(doc["path"]).write_text(doc["content"], encoding="utf-8")
    return {"opened": True, "path": doc["path"], "driver": driver, "created": True}


def document_save(payload: dict[str, Any], context: dict[str, Any]) -> dict[str, Any]:
    content = payload.get("content")
    path = payload.get("path")
    driver = _driver(context)
    if context.get("dry_run"):
        return {"dry_run": True, "driver": driver, "path": path}
    doc = _doc_state(context)
    if content is not None:
        doc["content"] = str(content)
    target = Path(str(path or doc.get("path") or _output_dir(context) / f"{doc.get('title', 'document')}.txt"))
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(doc.get("content") or "", encoding="utf-8")
    doc["path"] = str(target)
    return {"saved": True, "path": str(target), "bytes": target.stat().st_size, "driver": driver}


def writer_render(payload: dict[str, Any], context: dict[str, Any]) -> dict[str, Any]:
    text = str(payload.get("text") or payload.get("content") or "").strip()
    if not text:
        return {"ok": False, "error": "payload.text or payload.content is required"}
    title = str(payload.get("title") or "writer-output")
    fmt = str(payload.get("format") or "txt").lower()
    driver = _driver(context)
    if context.get("dry_run"):
        return {"dry_run": True, "driver": driver, "title": title, "format": fmt, "chars": len(text)}
    out = _output_dir(context)
    doc = _doc_state(context)
    if fmt == "docx":
        path = out / f"{title}.docx"
        _write_minimal_docx(path, text, title=title)
    else:
        path = out / f"{title}.txt"
        path.write_text(text, encoding="utf-8")
    doc.update({"path": str(path), "content": text, "title": title})
    return {"rendered": True, "path": str(path), "format": fmt, "chars": len(text), "driver": driver}


def _write_minimal_docx(path: Path, text: str, *, title: str) -> None:
    try:
        from docx import Document  # type: ignore

        document = Document()
        document.add_heading(title, level=1)
        for paragraph in text.split("\n"):
            document.add_paragraph(paragraph)
        document.save(path)
        return
    except Exception:
        pass
    path.write_text(text, encoding="utf-8")


def document_export_pdf(payload: dict[str, Any], context: dict[str, Any]) -> dict[str, Any]:
    source = payload.get("path")
    driver = _driver(context)
    doc = _doc_state(context)
    src_path = Path(str(source or doc.get("path") or ""))
    if context.get("dry_run"):
        pdf_path = str(src_path.with_suffix(".pdf")) if src_path.name else "document.pdf"
        return {"dry_run": True, "driver": driver, "pdf_path": pdf_path}
    if not src_path.exists():
        return {"ok": False, "error": f"source not found: {src_path}"}
    pdf_path = src_path.with_suffix(".pdf")
    if driver == "libreoffice" and _real_allowed(context) and shutil.which("soffice"):
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(src_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    else:
        pdf_path.write_text(f"PDF mock export of {src_path.name}\n", encoding="utf-8")
    return {"exported": True, "pdf_path": str(pdf_path), "source": str(src_path), "driver": driver}
